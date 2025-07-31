from docx import Document
import os
from datetime import datetime


def generate_docx(content: str) -> str:
    """
    Generates a Word document from the given content and returns the file path.

    Args:
        content: The text content to include in the Word document.

    Returns:
        A string with the file path where the Word document was saved.
    """
    doc = Document()
    doc.add_heading("Generated Report", level=1)
    doc.add_paragraph(content)

    output_dir = "data/word_outputs"
    os.makedirs(output_dir, exist_ok=True)

    filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    file_path = os.path.join(output_dir, filename)

    doc.save(file_path)

    return file_path
